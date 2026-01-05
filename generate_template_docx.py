from docx import Document
from docx.shared import Pt, RGBColor
import os

def create_template():
    doc = Document()

    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Title
    head = doc.add_heading('Πρότυπο Υποβολής Διαγωνίσματος', 0)
    head.alignment = 1

    # Metadata Section
    doc.add_heading('Στοιχεία Διαγωνίσματος (Συμπληρώστε τα):', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Μάθημα: Μαθηματικά')
    run.bold = True
    p.add_run(' (Αλλάξτε το σε Φυσική, Ιστορία, κλπ)')

    p = doc.add_paragraph()
    run = p.add_run('Τάξη: Β Γυμνασίου')
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Δυσκολία: Μέτριο')
    run.bold = True
    p.add_run(' (Επιλογές: Εύκολο, Μέτριο, Δύσκολο)')

    doc.add_paragraph('--------------------------------------------------')

    # Questions Section
    doc.add_heading('Ερωτήσεις', level=1)
    doc.add_paragraph('Οδηγίες: Γράψτε "ΕΡΩΤΗΣΗ" και τον αριθμό, μετά την εκφώνηση. Από κάτω "ΑΠΑΝΤΗΣΗ" και τη λύση. Αν θέλετε εικόνα, βάλτε την κάτω από την ερώτηση.')

    # Q1
    doc.add_heading('ΕΡΩΤΗΣΗ 1', level=2)
    doc.add_paragraph('Πόσο κάνει 5 + 5;')
    doc.add_paragraph('(Μπορείτε να εισάγετε εικόνα εδώ)')

    doc.add_heading('ΑΠΑΝΤΗΣΗ', level=3)
    doc.add_paragraph('Κάνει 10.')

    # Q2
    doc.add_heading('ΕΡΩΤΗΣΗ 2', level=2)
    doc.add_paragraph('Ποια είναι η πρωτεύουσα της Ελλάδας;')

    doc.add_heading('ΑΠΑΝΤΗΣΗ', level=3)
    doc.add_paragraph('Η Αθήνα.')

    # Save
    filename = "Template_Exam.docx"
    doc.save(filename)
    print(f"Created template: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_template()
