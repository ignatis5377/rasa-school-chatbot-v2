from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import os

# Create a dummy image
img_path = "sample_shape.png"
img = Image.new('RGB', (200, 200), color = 'red')
d = ImageDraw.Draw(img)
d.text((10,10), "Triangle ABC", fill=(255,255,0))
img.save(img_path)

# Create Word Doc
doc = Document()
doc.add_heading('Math Exam Questions', 0)

# Q1
doc.add_paragraph('Question: Υπολογίστε το εμβαδόν του παρακάτω σχήματος.')
doc.add_picture(img_path, width=Inches(2.0))
doc.add_paragraph('Answer: Το εμβαδόν είναι 200 τ.μ.')

# Q2
doc.add_paragraph('Question: Πόσο κάνει 10 + 10;')
doc.add_paragraph('Answer: Κάνει 20.')

filename = "Math_B_Easy_Sample.docx"
doc.save(filename)

print(f"Created {filename} successfully.")
print(f"ABS_PATH: {os.path.abspath(filename)}")
